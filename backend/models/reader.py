"""数据读取模块：DataReader 及子类。

封装不同类型文档的读取逻辑，体现 OOP 多态。
通过 reader_factory() 工厂方法按文件名/类型分发。
"""

from __future__ import annotations

import os
from abc import ABC, abstractmethod
from typing import IO, Union


class DataReader(ABC):
    """数据读取器抽象基类。

    所有具体 Reader 继承此类，实现 read() 方法。
    """

    @abstractmethod
    def read(self, source: Union[str, IO]) -> str:
        """读取数据源，返回纯文本。

        Args:
            source: 文件路径（str）或文件对象（IO）

        Returns:
            文本内容
        """
        ...

    def read_file(self, path: str) -> str:
        """便捷方法：按路径读取文件。"""
        with open(path, "rb") as f:
            return self.read(f)


class TextReader(DataReader):
    """纯文本读取器（.txt）。"""

    def read(self, source: Union[str, IO]) -> str:
        if isinstance(source, str):
            with open(source, "r", encoding="utf-8") as f:
                return f.read()
        # 文件对象
        source.seek(0)
        data = source.read()
        if isinstance(data, bytes):
            return data.decode("utf-8", errors="ignore")
        return data


class DocxReader(DataReader):
    """Word 文档读取器（.docx）。

    使用 python-docx 库（需 pip install python-docx）。
    若未安装，回退为纯文本读取（容错）。
    """

    def read(self, source: Union[str, IO]) -> str:
        try:
            from docx import Document
        except ImportError:
            # 未安装 python-docx，回退为文本读取
            fallback = TextReader()
            return f"[python-docx 未安装，按文本读取]\n{fallback.read(source)}"

        if isinstance(source, str):
            doc = Document(source)
        else:
            source.seek(0)
            doc = Document(source)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)


class ContractReader(DataReader):
    """合同读取器。

    读取合同文件（docx/txt），并标注文档类型。
    """

    def read(self, source: Union[str, IO]) -> str:
        # 根据扩展名分发
        content = _read_by_extension(source)
        return f"【合同】\n{content}"


class LetterReader(DataReader):
    """对外函件读取器。"""

    def read(self, source: Union[str, IO]) -> str:
        content = _read_by_extension(source)
        return f"【对外函件】\n{content}"


class TopicReader(DataReader):
    """上会议题读取器。"""

    def read(self, source: Union[str, IO]) -> str:
        content = _read_by_extension(source)
        return f"【上会议题】\n{content}"


class ProcurementReader(DataReader):
    """采购文件读取器。"""

    def read(self, source: Union[str, IO]) -> str:
        content = _read_by_extension(source)
        return f"【采购文件】\n{content}"


# ---------- 内部工具 ----------

def _read_by_extension(source: Union[str, IO]) -> str:
    """按扩展名选择读取器。"""
    # 文件路径
    if isinstance(source, str):
        ext = os.path.splitext(source)[1].lower()
        if ext == ".docx":
            return DocxReader().read(source)
        return TextReader().read(source)
    # 文件对象：尝试从文件名判断（Flask upload 的文件对象有 filename 属性）
    filename = getattr(source, "filename", "")
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".docx":
        return DocxReader().read(source)
    return TextReader().read(source)


# ---------- 工厂 ----------

def reader_factory(filename: str) -> DataReader:
    """根据文件名返回对应的 Reader 实例。

    Args:
        filename: 文件名（含中文标识或英文关键词）

    Returns:
        对应类型的 DataReader 子类实例
    """
    name = filename.lower()

    # 按文件名关键词匹配类型
    if any(kw in filename for kw in ["合同", "协议"]) or \
       "contract" in name:
        return ContractReader()
    if any(kw in filename for kw in ["函", "函件"]) or \
       "letter" in name:
        return LetterReader()
    if any(kw in filename for kw in ["议题", "上会"]) or \
       "topic" in name:
        return TopicReader()
    if any(kw in filename for kw in ["采购", "招标"]) or \
       "procure" in name:
        return ProcurementReader()

    # 默认：纯文本读取器
    return TextReader()
